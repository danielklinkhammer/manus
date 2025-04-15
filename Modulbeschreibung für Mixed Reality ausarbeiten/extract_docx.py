import docx
import os

def extract_text_from_docx(file_path):
    doc = docx.Document(file_path)
    full_text = []
    
    # Extract text from paragraphs
    for para in doc.paragraphs:
        full_text.append(para.text)
    
    # Extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                row_text.append(cell.text)
            full_text.append(" | ".join(row_text))
    
    return "\n".join(full_text)

if __name__ == "__main__":
    file_path = "/home/ubuntu/upload/Modulbeschreibung_MixedReality.docx"
    output_path = "/home/ubuntu/module_content.txt"
    
    if os.path.exists(file_path):
        text = extract_text_from_docx(file_path)
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(text)
        print(f"Content extracted and saved to {output_path}")
    else:
        print(f"File not found: {file_path}")
